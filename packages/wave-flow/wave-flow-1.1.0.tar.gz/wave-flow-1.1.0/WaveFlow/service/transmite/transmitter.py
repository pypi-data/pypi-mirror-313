import docx
import openpyxl


class Transmitter:
    """
    Scrolls through the provided files and finds the keys.
    Sequentially creates a .xlsx file with those keys as columns.
    """

    def __init__(self, docFiles: list[str], delimiter: str):
        self.__docFiles = docFiles
        self.__delimiter = delimiter
        self.__keys_from_paragraphs = set()
        self.__keys_from_tables = set()

        self.__extract_keys()

    def __extract_keys(self):
        for docFile in self.__docFiles:
            doc = docx.Document(docFile)

            self.__process_paragraphs(doc.paragraphs)

            for table in doc.tables:
                self.__process_table(table)

    def __process_paragraphs(self, paragraphs):
        for para in paragraphs:
            self.__extract_from_text(para.text, self.__keys_from_paragraphs)

    def __process_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                self.__extract_from_text(cell.text, self.__keys_from_tables)

    def __extract_from_text(self, text, keys_set):
        start_pos = 0
        while True:
            start_pos = text.find(self.__delimiter, start_pos)
            if start_pos == -1:
                break
            end_pos = text.find(self.__delimiter, start_pos + len(self.__delimiter))
            if end_pos == -1:
                break
            key = text[start_pos + len(self.__delimiter):end_pos]
            keys_set.add(key)
            start_pos = end_pos + len(self.__delimiter)

    def export(self, outputFile):
        if self.__keys_from_paragraphs or self.__keys_from_tables:
            workbook = openpyxl.Workbook()
            sheet = workbook.active
            sheet.title = "Extracted Keys"


            allKeys = list(self.__keys_from_paragraphs) + list(self.__keys_from_tables)
            sheet.append(allKeys)
            workbook.save(outputFile)
